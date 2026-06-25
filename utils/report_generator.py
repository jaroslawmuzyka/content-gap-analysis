import os
from io import BytesIO
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_html_report(sheets_dict, global_stats):
    html = """
    <!DOCTYPE html>
    <html lang="pl">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Content Gap Analysis - Raport</title>
        <style>
            :root {
                --primary: #4F46E5;
                --primary-dark: #4338CA;
                --bg: #F9FAFB;
                --surface: #FFFFFF;
                --text: #111827;
                --text-muted: #6B7280;
                --border: #E5E7EB;
                --success: #10B981;
                --danger: #EF4444;
            }
            body {
                font-family: 'Inter', system-ui, -apple-system, sans-serif;
                background-color: var(--bg);
                color: var(--text);
                margin: 0;
                padding: 0;
                line-height: 1.6;
            }
            .container {
                max-width: 1200px;
                margin: 0 auto;
                padding: 2rem;
            }
            .header {
                background: linear-gradient(135deg, var(--primary), var(--primary-dark));
                color: white;
                padding: 3rem 2rem;
                border-radius: 1rem;
                margin-bottom: 2rem;
                box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
            }
            .header h1 {
                margin: 0 0 0.5rem 0;
                font-size: 2.5rem;
            }
            .header p {
                margin: 0;
                opacity: 0.9;
                font-size: 1.2rem;
            }
            .stats-grid {
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
                gap: 1.5rem;
                margin-bottom: 3rem;
            }
            .stat-card {
                background: var(--surface);
                padding: 1.5rem;
                border-radius: 0.75rem;
                border: 1px solid var(--border);
                box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
                text-align: center;
                transition: transform 0.2s ease;
            }
            .stat-card:hover {
                transform: translateY(-5px);
            }
            .stat-value {
                font-size: 2.5rem;
                font-weight: 700;
                color: var(--primary);
                margin-bottom: 0.5rem;
            }
            .stat-label {
                font-size: 0.9rem;
                color: var(--text-muted);
                text-transform: uppercase;
                letter-spacing: 0.05em;
                font-weight: 600;
            }
            .section {
                background: var(--surface);
                padding: 2rem;
                border-radius: 1rem;
                border: 1px solid var(--border);
                margin-bottom: 2rem;
                box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
            }
            h2 {
                color: var(--text);
                border-bottom: 2px solid var(--border);
                padding-bottom: 0.75rem;
                margin-top: 0;
                margin-bottom: 1.5rem;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin-top: 1rem;
            }
            th, td {
                padding: 0.75rem 1rem;
                text-align: left;
                border-bottom: 1px solid var(--border);
            }
            th {
                background-color: #F3F4F6;
                font-weight: 600;
                color: var(--text-muted);
                text-transform: uppercase;
                font-size: 0.8rem;
                letter-spacing: 0.05em;
            }
            tr:hover td {
                background-color: #F9FAFB;
            }
            .badge {
                padding: 0.25rem 0.75rem;
                border-radius: 9999px;
                font-size: 0.85rem;
                font-weight: 500;
                display: inline-block;
            }
            .badge-success { background: #D1FAE5; color: #065F46; }
            .badge-danger { background: #FEE2E2; color: #991B1B; }
            .badge-warning { background: #FEF3C7; color: #92400E; }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>Raport Content Gap Analysis</h1>
                <p>Automatycznie wygenerowane zestawienie dla Twojego projektu</p>
            </div>
            
            <div class="stats-grid">
    """
    
    # Adding global stats
    for key, label in [
        ("przeanalizowane_frazy", "Frazy"),
        ("strony_konkurencji", "Strony Konkurencji"),
        ("strony_wlasne", "Strony Własne"),
        ("wygenerowane_pomysly", "Klastry/Pomysły"),
        ("content_gaps", "Luki Treści (Gaps)")
    ]:
        val = global_stats.get(key, 0)
        html += f"""
        <div class="stat-card">
            <div class="stat-value">{val}</div>
            <div class="stat-label">{label}</div>
        </div>
        """
        
    html += "</div>"
    
    # Adding Sections based on sheets
    # Focus on key sheets to not overwhelm HTML
    key_sheets = ["7. Weryfikacja Gap", "6. Brand Klastry", "8. Audyt Contentu", "4. Content Gap"]
    
    for sheet_name in key_sheets:
        if sheet_name in sheets_dict and not sheets_dict[sheet_name].empty:
            df = sheets_dict[sheet_name]
            html += f'<div class="section"><h2>{sheet_name}</h2>'
            
            # Format dataframe to html
            # Limit rows in HTML for performance if very large, maybe top 100
            display_df = df.head(100)
            
            # Basic HTML table generation
            html += "<table><thead><tr>"
            for col in display_df.columns:
                html += f"<th>{col}</th>"
            html += "</tr></thead><tbody>"
            
            for _, row in display_df.iterrows():
                html += "<tr>"
                for col in display_df.columns:
                    val = str(row[col]) if pd.notnull(row[col]) else ""
                    # Simple color coding for known columns
                    if col in ["AI Verdict", "Rekomendacja", "Priorytet", "Weryfikacja"]:
                        if "pasuje" in val.lower() or "wysoki" in val.lower() or "dobrze" in val.lower() or "brak_treści" in val.lower():
                            html += f'<td><span class="badge badge-success">{val}</span></td>'
                        elif "nie pasuje" in val.lower() or "niski" in val.lower() or "błąd" in val.lower() or "pokrywa_się" in val.lower():
                            html += f'<td><span class="badge badge-danger">{val}</span></td>'
                        else:
                            html += f'<td><span class="badge badge-warning">{val}</span></td>'
                    else:
                        html += f"<td>{val}</td>"
                html += "</tr>"
            
            html += "</tbody></table>"
            if len(df) > 100:
                html += f"<p style='text-align:center; color:var(--text-muted); margin-top:1rem;'><i>Wyświetlono pierwsze 100 rekordów z {len(df)}</i></p>"
            html += "</div>"
            
    html += """
        </div>
    </body>
    </html>
    """
    
    return html.encode('utf-8')

def generate_docx_report(sheets_dict, global_stats):
    doc = Document()
    
    # Title
    doc.add_heading('Raport Content Gap Analysis', 0)
    
    # Global Stats
    doc.add_heading('1. Podsumowanie Analizy', level=1)
    p = doc.add_paragraph()
    p.add_run('Kluczowe metryki z przeprowadzonej analizy:\n').bold = True
    p.add_run(f"- Przeanalizowane frazy: {global_stats.get('przeanalizowane_frazy', 0)}\n")
    p.add_run(f"- Strony konkurencji: {global_stats.get('strony_konkurencji', 0)}\n")
    p.add_run(f"- Analizowane strony własne: {global_stats.get('strony_wlasne', 0)}\n")
    p.add_run(f"- Wygenerowane Klastry/Pomysły: {global_stats.get('wygenerowane_pomysly', 0)}\n")
    p.add_run(f"- Zidentyfikowane Luki (Content Gaps): {global_stats.get('content_gaps', 0)}\n")
    
    # Section: Content Gaps
    if "7. Weryfikacja Gap" in sheets_dict and not sheets_dict["7. Weryfikacja Gap"].empty:
        df_ver = sheets_dict["7. Weryfikacja Gap"]
        gaps = df_ver[df_ver["Weryfikacja"].astype(str).str.contains("BRAK_TRE", case=False, na=False)]
        
        doc.add_heading('2. Najważniejsze Luki Treści (Zupełnie nowe tematy do wdrożenia)', level=1)
        doc.add_paragraph(f"Zidentyfikowano {len(gaps)} pomysłów, które nie są obecnie pokryte na Twojej stronie i stanowią doskonałą okazję do wdrożenia.")
        
        # Display top 10
        for idx, row in gaps.head(10).iterrows():
            p = doc.add_paragraph(style='List Bullet')
            prod = str(row.get("Recommended Product", ""))
            reason = str(row.get("Reasoning", ""))
            p.add_run(f"Produkt: {prod}").bold = True
            p.add_run(f"\nUzasadnienie: {reason}")
            
    # Section: Brand Clusters
    if "6. Brand Klastry" in sheets_dict and not sheets_dict["6. Brand Klastry"].empty:
        df_brand = sheets_dict["6. Brand Klastry"]
        doc.add_heading('3. Wnioski z Fraz Brandowych (Klastry)', level=1)
        doc.add_paragraph(f"Wygenerowano {len(df_brand)} klastrów zapytań związanych bezpośrednio z marką.")
        
        # Top 10 High Priority
        high_pri = df_brand[df_brand["Priorytet"].astype(str).str.contains("Wysoki", case=False, na=False)]
        for idx, row in high_pri.head(10).iterrows():
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{str(row.get('Nazwa Klastra', ''))}").bold = True
            p.add_run(f" (URL Docelowy: {str(row.get('Adres URL (Docelowy / Produkt)', ''))})\n")
            p.add_run(f"Zalecenie: {str(row.get('Rekomendowana Akcja', ''))}")
            
    # Section: Audyt
    if "8. Audyt Contentu" in sheets_dict and not sheets_dict["8. Audyt Contentu"].empty:
        df_audit = sheets_dict["8. Audyt Contentu"]
        doc.add_heading('4. Podsumowanie Audytu Contentu', level=1)
        doc.add_paragraph("Analiza jakości stron i ich przygotowania dla sztucznej inteligencji (AI Readiness).")
        
        for idx, row in df_audit.iterrows():
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"URL: {str(row.get('URL', ''))}\n").bold = True
            p.add_run(f"Ogólna ocena: {str(row.get('Ocena ogólna (1-10)', ''))}\n")
            p.add_run(f"Główny wniosek: {str(row.get('Główny wniosek / Największy problem', ''))}")

    # Save
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
